"""產出引擎：複製前月 10 份範本 → 重新命名 → 逐檔更新日期與數據。"""

from __future__ import annotations

import glob
import os
import re
import shutil
from datetime import date

from docx import Document

from .dates import Meta, build_meta
from .docx_utils import (get_cell_text, iter_paragraphs, replace_text_in_doc,
                          set_cell_text, set_paragraph_text, update_asof_date,
                          update_dates)
from .parse_excel import MinistryExcel

DOC_SUFFIXES = [
    "-1議程", "-2事工成果統計表", "-3收入支用統計表", "-4各項奉獻",
    "-5贈經事工(除學校)", "-6學校贈經統計表", "-7會員及地界教會代禱項目",
    "-8早禱會及月例會輪值表", "-9年度教會見證統計表(橫式)", "-10會員名冊(橫式)",
]


# 下次月例會：日期一律由 meta.next_meeting_date（下個月第四個禮拜天）推導，
# 地點原則上固定，但年會月份會變（如 2026/08 是「年會場地」），所以介面可覆寫。
DEFAULT_NEXT_VENUE = "土城清水教會"
DEFAULT_NEXT_TIME = "下午4:00"
_RE_NEXT_MEETING = re.compile(r"^(\s*[一二三四五六七八九十]+、\s*下次月例會預定)")


def doc_filename(roc: int, month: int, suffix: str) -> str:
    return f"月例會議程{roc}年{month}月{suffix}.docx"


def file_number(fname: str) -> int:
    m = re.search(r"月\s*-(\d+)", fname) or re.search(r"-(\d+)", fname)
    return int(m.group(1)) if m else 0


# ── Step 1：初始化當月工作區 ─────────────────────────────
def init_month(base_dir: str, meta: Meta, template_dir: str | None = None) -> dict:
    """建立當月資料夾、由前月複製 10 份 docx 並重新命名。"""
    work_dir = os.path.join(base_dir, meta.work_dir_name)
    prev_dir = template_dir or os.path.join(base_dir, f"{meta.prev_year}年{meta.prev_month}月月例會")

    if not os.path.isdir(prev_dir):
        return {"ok": False, "error": f"找不到範本資料夾：{prev_dir}", "work_dir": work_dir}

    os.makedirs(work_dir, exist_ok=True)
    os.makedirs(os.path.join(work_dir, "_inputs"), exist_ok=True)

    files, skipped = [], []
    for suffix in DOC_SUFFIXES:
        src_name = doc_filename(meta.prev_roc_year, meta.prev_month, suffix)
        src = os.path.join(prev_dir, src_name)
        if not os.path.exists(src):
            cands = glob.glob(os.path.join(prev_dir, f"*{glob.escape(suffix)}.docx"))
            if not cands:
                skipped.append(suffix)
                continue
            src = cands[0]
        dst = os.path.join(work_dir, doc_filename(meta.roc_year, meta.report_month, suffix))
        if not os.path.exists(dst):
            shutil.copy2(src, dst)
        files.append({"suffix": suffix, "name": os.path.basename(dst),
                      "path": dst, "status": "copied"})

    # 附件（如 -A.行事曆2026-2027北區.docx）一併帶過來
    prev_prefix = f"月例會議程{meta.prev_roc_year}年{meta.prev_month}月"
    for src in glob.glob(os.path.join(prev_dir, f"{prev_prefix}-A*.docx")):
        tail = os.path.basename(src)[len(prev_prefix):]
        dst = os.path.join(work_dir,
                           f"月例會議程{meta.roc_year}年{meta.report_month}月{tail}")
        if not os.path.exists(dst):
            shutil.copy2(src, dst)
        files.append({"suffix": tail.rstrip(".docx"), "name": os.path.basename(dst),
                      "path": dst, "status": "copied", "extra": True})

    return {"ok": True, "work_dir": work_dir, "template_dir": prev_dir,
            "files": files, "missing": skipped}


# ── Step 2：逐檔更新 ─────────────────────────────────────
def generate_all(work_dir: str, meta: Meta, excel_path: str | None = None,
                 model: dict | None = None, as_of: date | None = None) -> dict:
    """更新工作區內的 10 份 docx。回傳每份的狀態與變更紀錄。

    `as_of` 是製表日（各文件標題下方那個日期），預設今天。雲端跑在 UTC，
    台灣時間深夜產報告會差一天，所以介面會把使用者當地的日期送進來。
    """
    model = model or {}
    as_of = as_of or date.today()
    data = None
    if excel_path and os.path.exists(excel_path):
        data = MinistryExcel(excel_path)

    prefix = f"月例會議程{meta.roc_year}年{meta.report_month}月"
    paths = sorted(glob.glob(os.path.join(work_dir, f"{prefix}-*.docx")),
                   key=lambda p: file_number(os.path.basename(p)))
    if not paths:
        return {"ok": False, "error": f"工作區找不到 {prefix}-*.docx", "results": []}

    results = []
    for path in paths:
        fname = os.path.basename(path)
        num = file_number(fname)
        doc = Document(path)
        log = update_dates(doc, meta)
        log += update_asof_date(doc, as_of)
        status, notes = "ok", []

        # API 資料優先於 Excel（新財年 Excel 產不出來，API 才有正確數字）
        api_stats = (model.get("ministry_stats_api") or None)
        api_churches = (model.get("church_testimony_api") or None)
        src = (model.get("_sources") or [""])[-1]

        try:
            if num == 1:
                notes += _update_agenda(doc, meta, model)
                notes += _update_next_meeting(doc, meta, model)
            elif num == 2:
                if api_stats:
                    notes += _update_ministry_from_api(doc, api_stats, src)
                elif data:
                    notes += _update_ministry(doc, data)
            elif num == 4 and data:
                notes += _update_offerings(doc, data)
            elif num == 5:
                notes += _update_bible_giving(doc, meta, model)
            elif num == 6:
                notes += _update_schools(doc, meta, model)
            elif num == 9:
                if api_churches:
                    notes += _update_church_testimony_api(doc, api_churches, api_stats)
                elif data:
                    notes += _update_church_testimony(doc, data)
        except Exception as exc:  # 不讓單一檔失敗中斷整批
            status, notes = "warn", notes + [f"更新時發生例外：{exc}"]

        doc.save(path)
        results.append({
            "file": fname, "num": num, "path": path, "status": status,
            "date_changes": log, "notes": notes,
            "data_updated": bool(notes),
        })

    return {"ok": True, "work_dir": work_dir, "results": results,
            "as_of": as_of.isoformat(),
            "excel": os.path.basename(excel_path) if excel_path else None}


# ── 各文件更新 ───────────────────────────────────────────
def _find_row(table, keyword: str, col: int = 0):
    for row in table.rows:
        if keyword in get_cell_text(row.cells[col]):
            return row
    return None


def _update_next_meeting(doc, meta: Meta, model: dict) -> list[str]:
    """-1 議程「八、下次月例會預定 …」整行重寫。

    這一行原本沒有任何更新邏輯，只能靠 update_dates() 碰運氣——而它寫的是
    零補位的 `2026/07/26`，跟 update_dates() 找的 `2026/7/` 對不上，所以整行
    會原封不動留著上個月的內容。7 月那份還留著年會的「地點:年會場地」，
    直接複製就會變成下個月也在年會場地開會。
    """
    nm = model.get("next_meeting") or {}
    d = date.fromisoformat(meta.next_meeting_date)
    text = (nm.get("date_text") or "").strip() or \
        f"{d.year}/{d.month:02d}/{d.day:02d}{DEFAULT_NEXT_TIME}"
    venue = (nm.get("venue") or "").strip() or DEFAULT_NEXT_VENUE

    for para in iter_paragraphs(doc):
        m = _RE_NEXT_MEETING.match(para.text)
        if not m:
            continue
        new = f"{m.group(1)} {text} 地點:{venue}"
        if para.text.strip() == new.strip():
            return []
        set_paragraph_text(para, new)
        return [f"下次月例會 → {text}　地點:{venue}"]
    return ["⚠️ 找不到「下次月例會預定」段落，未更新"]


def _update_agenda(doc, meta: Meta, model: dict) -> list[str]:
    """-1 議程：年度主題／目標；未定項目標 TODO。"""
    notes = []
    annual = model.get("annual") or {}
    if theme := annual.get("theme"):
        if replace_text_in_doc(doc, "【年度主題】", theme):
            notes.append(f"年度主題 → {theme}")
    if bt := annual.get("bible_target"):
        notes.append(f"年度贈經目標 {bt:,} 本（供人工核對）")
    todos = model.get("agenda_todo", {})
    for key, label in (("motions", "臨時動議"), ("guest_reports", "來賓報告"),
                       ("next_events", "下月活動預告")):
        if items := todos.get(key):
            notes.append(f"{label}：{len(items)} 筆（【TODO】待人工置入）")
    return notes


# -2 docx 欄索引 → Excel 匯總欄語意名稱（已用 2026/06 匯出檔驗證）
MINISTRY_COL_MAP = {
    1: ("members_brother", "人"),   # 會員數 弟兄
    2: ("members_sister", "人"),    # 會員數 姐妹
    3: ("fee_brother", "人"),       # 會費 弟兄
    4: ("fee_sister", "人"),        # 會費 姐妹
    5: ("scripture_total", ""),     # 贈送聖經(本)
    6: ("church_count", "次"),      # 教會見證 次數
    7: ("church_offering", ""),     # 教會聖奉
    8: ("bible_brother", ""),       # 會員聖奉 弟兄
    9: ("bible_sister", ""),        # 會員聖奉 姊妹
    10: ("scripture_sister", ""),   # 姊妹贈經
}

MINISTRY_ROW_LABELS = {"目標": "target", "成果": "actual",
                       "差額": "diff", "達成率": "rate"}


def _fmt_stat(key: str, val, unit: str) -> str:
    if val in (None, ""):
        return "-"
    if key == "rate":
        # Excel 的達成率以小數儲存（1 = 100%）
        return f"{val * 100:.0f}%" if isinstance(val, (int, float)) else f"{val}"
    if isinstance(val, (int, float)):
        return f"{int(round(val)):,}{unit}"
    return f"{val}{unit}"


def _update_ministry_from_api(doc, stats: dict, source: str = "") -> list[str]:
    """-2 事工成果統計表：以 Grafana API 的 category 資料填四列。

    新財年目標未設定時 goal=None，目標／差額／達成率填「-」，
    與實際 6 月報告的呈現一致。
    """
    from .grafana import MINISTRY_CATEGORIES
    table = doc.tables[0]
    notes, filled, missing = [], 0, []

    # docx 列標籤 → API 欄位名（與 Excel 版的 target/actual/… 不同）
    api_fields = {"目標": "goal", "成果": "value", "差額": "diff", "達成率": "rate"}

    for label, key in api_fields.items():
        row = _find_row(table, label)
        if row is None:
            notes.append(f"找不到「{label}」列")
            continue
        for col, category in MINISTRY_CATEGORIES.items():
            if col >= len(row.cells):
                continue
            rec = stats.get(category)
            if rec is None:
                if label == "成果" and category not in missing:
                    missing.append(category)
                set_cell_text(row.cells[col], "-")
                continue
            unit = "人" if category in ("增加弟兄", "增加姊妹", "弟兄會費",
                                        "姊妹會費") else (
                   "次" if category == "教會見證" else "")
            set_cell_text(row.cells[col], _fmt_stat(key, rec[key], unit))
            filled += 1

    act = {c: (stats.get(c) or {}).get("value") for c in
           ("弟兄會費", "姊妹會費", "教會聖奉", "贈送聖經")}
    notes.append(f"四列共 {filled} 欄已套用{('（來源：' + source + '）') if source else ''}"
                 f"（成果：會費 {act['弟兄會費']}/{act['姊妹會費']}、"
                 f"教會聖奉 {act['教會聖奉']}、贈經 {act['贈送聖經']}）")
    if stats and any(v["goal"] is not None for v in stats.values()):
        notes.append("✅ 年度目標已帶入，差額與達成率同步更新")
    if missing:
        notes.append(f"下列項目本財年查無資料，已填「-」：{'、'.join(missing)}")
    if stats and all(v["goal"] is None for v in stats.values()):
        notes.append("⚠️ 本財年年度目標尚未設定，目標／差額／達成率皆為「-」")
    return notes


def _update_ministry(doc, data: MinistryExcel) -> list[str]:
    """-2 事工成果統計表：目標／成果／差額／達成率四列直接對映 Excel Row 5-8。"""
    table = doc.tables[0]
    notes, filled = [], 0

    for label, key in MINISTRY_ROW_LABELS.items():
        row = _find_row(table, label)
        if row is None:
            notes.append(f"找不到「{label}」列")
            continue
        stats = data.summary[key]
        for col, (field, unit) in MINISTRY_COL_MAP.items():
            if col >= len(row.cells):
                continue
            val = stats.get(field)
            set_cell_text(row.cells[col],
                          _fmt_stat(key, val, "" if key in ("rate",) else unit))
            filled += 1

    act = data.summary["actual"]
    notes.append(f"四列共 {filled} 欄已對映 Excel Row 5-8"
                 f"（成果：會費 {act['fee_brother']}/{act['fee_sister']}、"
                 f"贈經 {act['scripture_total']}、教會聖奉 {act['church_offering']}）")
    return notes


def _update_offerings(doc, data: MinistryExcel) -> list[str]:
    """-4 各項奉獻：依會員姓名比對，更新會費日期與聖經奉獻。"""
    table = doc.tables[0]
    updated = 0
    for row in table.rows[2:]:
        cells = row.cells
        if len(cells) < 7:
            continue
        label = get_cell_text(cells[0])
        if label in ("目標", "成果", "差額", "達成率", ""):
            continue
        for name_col, fee_col, bible_col, field in (
            (1, 3, 5, "brother"), (2, 4, 6, "sister"),
        ):
            name = get_cell_text(cells[name_col])
            if not name:
                continue
            m = data.member_by_name(name)
            if not m:
                continue
            fee = m[f"fee_{field}"]
            bible = m[f"bible_{field}"]
            if get_cell_text(cells[fee_col]) in ("安息", "退會"):  # 保留人工註記
                continue
            if fee:
                set_cell_text(cells[fee_col], str(fee))
                updated += 1
            if bible:
                set_cell_text(cells[bible_col], f"{int(bible):,}")
                updated += 1

    # 四列匯總直接取 Excel Row 5-8（會費人數、聖經奉獻）
    labels = 0
    for label, key in (("目標", "target"), ("成果", "actual"),
                       ("差額", "diff"), ("達成率", "rate")):
        row = _find_row(table, label)
        if row is None or len(row.cells) < 7:
            continue
        s = data.summary[key]
        for col, field in ((3, "fee_brother"), (4, "fee_sister"),
                           (5, "bible_brother"), (6, "bible_sister")):
            set_cell_text(row.cells[col], _fmt_stat(key, s.get(field), ""))
        labels += 1
    return [f"各項奉獻已更新 {updated} 個欄位、{labels} 列匯總已對映 Excel"]


_CHURCH_NOISE = ("基督教", "基督", "台灣", "臺灣", "教會", "禮拜堂", "福音中心",
                 "堂", "會", "（北中）", "(北中)", " ", "　")


def _church_key(name: str) -> str:
    """教會名正規化：去掉宗派／型態等雜訊字，只留核心地名。

    例：「板城靈糧堂教會」與「板城靈糧堂」→ 同一 key「板城靈糧」。
    """
    s = str(name or "")
    for w in _CHURCH_NOISE:
        s = s.replace(w, "")
    return s


def _longest_common(a: str, b: str) -> str:
    """最長共同子字串（用於教會名不同寫法的比對）。"""
    if not a or not b:
        return ""
    best, la, lb = "", len(a), len(b)
    prev = [0] * (lb + 1)
    for i in range(1, la + 1):
        cur = [0] * (lb + 1)
        for j in range(1, lb + 1):
            if a[i - 1] == b[j - 1]:
                cur[j] = prev[j - 1] + 1
                if cur[j] > len(best):
                    best = a[i - cur[j]:i]
        prev = cur
    return best


# 共同字串達此長度即視為同一間教會（如「愛加倍」）
MIN_COMMON = 3


def _match_church(cname: str, chs: list[dict]):
    """精確 → 正規化包含 → 最長共同子字串。回傳 (教會資料, 比對方式)。

    實際資料兩邊寫法常不同：
      「板城靈糧堂教會」vs「板城靈糧堂」→ 包含比對
      「樹林愛加倍教會」vs「愛加倍浸信會」→ 共同子字串「愛加倍」
    """
    for c in chs:
        if str(c["church"]) == cname:
            return c, "exact"

    k = _church_key(cname)
    if len(k) < 2:
        return None, ""

    for c in chs:
        ck = _church_key(c["church"])
        if ck and (k == ck or k in ck or ck in k):
            return c, "fuzzy"

    best, best_len = None, 0
    for c in chs:
        common = _longest_common(k, _church_key(c["church"]))
        if len(common) >= MIN_COMMON and len(common) > best_len:
            best, best_len = c, len(common)
    return (best, "partial") if best else (None, "")


def _update_church_testimony(doc, data: MinistryExcel) -> list[str]:
    """-9 年度教會見證統計表：各教會日期／講員／金額 + 最後匯總列（易遺漏）。"""
    table = doc.tables[0]
    chs = data.churches()
    used, exact, fuzzy, unmatched = set(), 0, 0, []

    for row in table.rows[2:]:
        cells = row.cells
        if len(cells) < 12:
            continue
        cname = get_cell_text(cells[1])
        if not cname or "小計" in cname or "目標" in cname:
            continue
        c, how = _match_church(cname, chs)
        if not c:
            unmatched.append(cname)
            continue
        if c["date"]:
            set_cell_text(cells[5], str(c["date"]).replace("-", "/"))
        if c["speaker"]:
            set_cell_text(cells[9], str(c["speaker"]))
        set_cell_text(cells[11], f"{int(c['amount']):,}" if c["amount"] else "")
        used.add(str(c["church"]))
        exact += how == "exact"
        fuzzy += how == "fuzzy"

    # ── 匯總列：一律採 Excel 官方匯總值，不自行加總 ────────
    act, tgt = data.summary["actual"], data.summary["target"]
    total_amount = int(act.get("church_offering") or 0)
    target_amount = int(tgt.get("church_offering") or 0)
    done_count = int(act.get("church_count") or 0)
    target_count = int(tgt.get("church_count") or 0)
    rate = f"{total_amount / target_amount * 100:.0f}%" if target_amount else "-"

    summary_notes = []
    for row in table.rows[-4:]:
        text = get_cell_text(row.cells[0])
        if "小計" in text and len(row.cells) > 11:
            set_cell_text(row.cells[11], f"{total_amount:,}")
            summary_notes.append(f"小計 {total_amount:,}")
        elif "目標間數" in text:
            set_cell_text(row.cells[0],
                          f"目標間數: {target_count or '-'}    已達成: {done_count} 間")
            summary_notes.append(f"間數 {done_count}/{target_count or '-'}")
        elif "目標金額" in text:
            set_cell_text(row.cells[0],
                          f"目標金額: {target_amount:,}    已達成: {total_amount:,} 元"
                          f"    達成率: {rate}")
            summary_notes.append(f"金額 {total_amount:,}（達成率 {rate}）")

    notes = [f"教會見證已更新 {exact + fuzzy} 間（精確 {exact}、模糊 {fuzzy}）",
             "匯總列：" + ("、".join(summary_notes) if summary_notes else "未找到，請人工檢查")]
    if unmatched:
        notes.append(f"⚠️ 未對應到 Excel 的教會 {len(unmatched)} 間："
                     + "、".join(unmatched[:8]))
    left = [c["church"] for c in chs if str(c["church"]) not in used]
    if left:
        notes.append(f"⚠️ Excel 有但表中未列 {len(left)} 間：" + "、".join(map(str, left[:8])))
    return notes


def _update_church_testimony_api(doc, churches: list[dict],
                                 stats: dict | None) -> list[str]:
    """-9 年度教會見證：以 API 資料更新，匯總列取 API 的教會見證/教會聖奉。"""
    table = doc.tables[0]
    used, exact, fuzzy, unmatched = set(), 0, 0, []

    for row in table.rows[2:]:
        cells = row.cells
        if len(cells) < 12:
            continue
        cname = get_cell_text(cells[1])
        if not cname or "小計" in cname or "目標" in cname:
            continue
        c, how = _match_church(cname, churches)
        if not c:
            unmatched.append(cname)
            continue
        if c["date"]:
            set_cell_text(cells[5], str(c["date"]).replace("-", "/"))
        if c["speaker"]:
            set_cell_text(cells[9], str(c["speaker"]))
        set_cell_text(cells[11], f"{int(c['amount']):,}" if c["amount"] else "")
        used.add(str(c["church"]))
        exact += how == "exact"
        fuzzy += how in ("fuzzy", "partial")

    cnt = (stats or {}).get("教會見證") or {}
    amt = (stats or {}).get("教會聖奉") or {}
    total = int(amt.get("value") or 0)
    tgt_amt = amt.get("goal")
    done = int(cnt.get("value") or len(churches))
    tgt_cnt = cnt.get("goal")
    rate = f"{total / tgt_amt * 100:.0f}%" if tgt_amt else "-"

    summary = []
    for row in table.rows[-4:]:
        text = get_cell_text(row.cells[0])
        if "小計" in text and len(row.cells) > 11:
            set_cell_text(row.cells[11], f"{total:,}")
            summary.append(f"小計 {total:,}")
        elif "目標間數" in text:
            set_cell_text(row.cells[0],
                          f"目標間數: {tgt_cnt if tgt_cnt is not None else '-'}"
                          f"    已達成: {done} 間")
            summary.append(f"間數 {done}/{tgt_cnt if tgt_cnt is not None else '-'}")
        elif "目標金額" in text:
            set_cell_text(row.cells[0],
                          f"目標金額: {f'{int(tgt_amt):,}' if tgt_amt else '-'}"
                          f"    已達成: {total:,} 元    達成率: {rate}")
            summary.append(f"金額 {total:,}（{rate}）")

    notes = [f"教會見證取自 API，已更新 {exact + fuzzy} 間（精確 {exact}、模糊 {fuzzy}）",
             "匯總列：" + ("、".join(summary) if summary else "未找到，請人工檢查")]
    if unmatched:
        notes.append(f"⚠️ 表中有但 API 無資料的教會 {len(unmatched)} 間："
                     + "、".join(unmatched[:8]))
    left = [c["church"] for c in churches if str(c["church"]) not in used]
    if left:
        notes.append(f"⚠️ API 有但表中未列 {len(left)} 間：" + "、".join(map(str, left[:8])))
    return notes


def _update_schools(doc, meta: Meta, model: dict) -> list[str]:
    """-6 學校贈經統計表：填入本月的學校場次。

    來源是年度贈經計畫（整個財年固定），非 LINE。空白列才填，
    已有內容的列一律保留——承辦人可能已手動補過實際本數與同工。
    """
    from .parse_plan import schools_of_month

    plan = model.get("distribution_plan") or {}
    if not plan.get("schools"):
        return []

    all_month = schools_of_month(plan, meta.report_year, meta.report_month,
                                 include_undated=True)
    this_month = [s for s in all_month if s["date"]]
    pending = [s for s in all_month if not s["date"]]
    if not all_month:
        return [f"年度計畫中 {meta.report_month} 月無學校場次"]
    if not this_month:
        return ["日期未定，未自動填入（請定案後手動補）："
                + "、".join(s["school"] for s in pending)]

    table = doc.tables[0]
    # 表頭：日期 / 學校、醫院、旅館 / 數量 / 参與同工 / 弟兄 / 姊妹
    existing = {get_cell_text(r.cells[1]) for r in table.rows[1:]
                if get_cell_text(r.cells[1])}

    filled, skipped = 0, []
    row_iter = (r for r in table.rows[1:] if not get_cell_text(r.cells[1]))
    for s in this_month:
        if s["school"] in existing:
            skipped.append(s["school"])
            continue
        row = next(row_iter, None)
        if row is None:
            skipped.append(s["school"] + "（表格列數不足）")
            continue
        date_txt = s["date"][5:].replace("-", "/") if s["date"] else s["date_raw"]
        set_cell_text(row.cells[0], date_txt)
        set_cell_text(row.cells[1], s["school"])
        if s["count"]:
            set_cell_text(row.cells[2], str(int(s["count"])))
        filled += 1

    notes = []
    if pending:
        notes.append("日期未定，未自動填入（請定案後手動補）："
                     + "、".join(s["school"] for s in pending))
    if filled:
        notes.append(f"已填入本月 {filled} 場學校贈經："
                     + "、".join(s["school"] for s in this_month)[:60])
    if skipped:
        notes.append(f"略過 {len(skipped)} 筆（表中已有或空列不足）："
                     + "、".join(skipped[:5]))
    for s in this_month:
        if s["status"] != "已排定":
            notes.append(f"⚠️ {s['school']}：{s['status']}（{s['date_raw'] or '無日期'}）")
    notes.append("數量與參與同工需人工補齊")
    return notes


def _update_bible_giving(doc, meta: Meta, model: dict) -> list[str]:
    """-5 贈經事工：列出當月排程（僅回報，內容由人確認後貼入）。"""
    sched = (model.get("bible_giving") or {}).get("schedule") or []
    this_month = [s for s in sched if s.get("date", "") and
                  s["date"][:7] == f"{meta.report_year}-{meta.report_month:02d}"]
    if not this_month:
        return []
    return [f"本月贈經排程 {len(this_month)} 筆："
            + "、".join(f"{s['date'][5:]} {s['target']}" for s in this_month)]


# ── 便利入口 ─────────────────────────────────────────────
def run(base_dir: str, year: int, month: int, meeting_date: str | None = None,
        excel_path: str | None = None, model: dict | None = None) -> dict:
    meta = build_meta(year, month, meeting_date)
    init = init_month(base_dir, meta)
    if not init["ok"]:
        return init
    gen = generate_all(init["work_dir"], meta, excel_path, model)
    gen["meta"] = meta.to_dict()
    gen["init"] = init
    return gen
