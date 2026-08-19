"""python-docx 共用工具：多段落 cell 寫入、跨 run 文字替換、日期替換。"""

from __future__ import annotations

import re
from datetime import date

from docx import Document

from .dates import chinese_month_name, roc_year

# 財年期間字樣，如 2025-2026 / 2025~2026
_RE_PERIOD = re.compile(r"20\d{2}\s*[-–~]\s*20\d{2}")
# 完整年月日，如 2026年7月24日
_RE_YMD = re.compile(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日")


def set_cell_text(cell, new_text) -> None:
    """清除 cell 全部段落的全部 run，於第一段寫入新文字。

    表格 cell 常含多個 paragraph（如「13人 / 36,000 / 元」為 3 段），
    只改第一個 run 會留下殘影，故需清乾淨再寫。
    """
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = str(new_text)
    else:
        cell.paragraphs[0].text = str(new_text)


def get_cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _replace_in_paragraph(para, old: str, new: str) -> int:
    """在段落內替換文字，支援文字被拆成多個 run 的情況。"""
    hits = 0
    # 先試單一 run（保留格式最佳）
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            hits += 1
    if hits or old not in para.text:
        return hits
    # 跨 run：合併到第一個 run，其餘清空
    full = para.text.replace(old, new)
    if para.runs:
        para.runs[0].text = full
        for run in para.runs[1:]:
            run.text = ""
        hits = 1
    return hits


def set_paragraph_text(para, text: str) -> None:
    """整段改寫：文字寫進第一個 run，其餘清空（保留段落格式）。

    跟 set_cell_text 同樣的道理——只改第一個 run 而不清掉其他 run 會留殘影。
    """
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.text = text


def iter_paragraphs(doc):
    """走訪文件內所有段落（含表格 cell、巢狀表格）。"""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
                for t in cell.tables:
                    for r in t.rows:
                        for c in r.cells:
                            for p in c.paragraphs:
                                yield p


def replace_text_in_doc(doc, old: str, new: str) -> int:
    """全文件替換，回傳命中次數。"""
    if old == new:
        return 0
    return sum(_replace_in_paragraph(p, old, new) for p in iter_paragraphs(doc))


def update_dates(doc, meta) -> list[str]:
    """依鐵則替換日期：先具體日期，後通用月份。回傳變更紀錄。"""
    log: list[str] = []
    prev_md = date.fromisoformat(meta.prev_meeting_date)
    curr_md = date.fromisoformat(meta.meeting_date)

    def rep(old, new, label):
        n = replace_text_in_doc(doc, old, new)
        if n:
            log.append(f"{label}: {old} → {new} ({n})")

    # ── 1. 具體日期（務必最先）──────────────────────────
    rep(f"{meta.prev_month}月{prev_md.day}日",
        f"{meta.report_month}月{curr_md.day}日", "會議日期")
    rep(f"{meta.prev_month}/{prev_md.day}",
        f"{meta.report_month}/{curr_md.day}", "會議日期(斜線)")
    rep(f"{meta.prev_year}年{meta.prev_month}月{prev_md.day}日",
        f"{meta.report_year}年{meta.report_month}月{curr_md.day}日", "完整日期")

    # ── 2. 中文月份 ────────────────────────────────────
    rep(chinese_month_name(meta.prev_month),
        chinese_month_name(meta.report_month), "中文月份")

    # ── 3. 通用年月 ────────────────────────────────────
    rep(f"{meta.prev_roc_year}年{meta.prev_month}月",
        f"{meta.roc_year}年{meta.report_month}月", "民國年月")
    rep(f"{meta.prev_year}年{meta.prev_month}月",
        f"{meta.report_year}年{meta.report_month}月", "西元年月")

    # ── 4. 斜線年月（如列印日期 2026/6/24，保留原日）────
    rep(f"{meta.prev_year}/{meta.prev_month}/",
        f"{meta.report_year}/{meta.report_month}/", "斜線年月")

    # ── 5. 財年期間：掃描 20xx-20xx，一律校正為當前財年 ────
    seen: set[str] = set()
    for para in iter_paragraphs(doc):
        for m in _RE_PERIOD.finditer(para.text):
            if m.group(0) != meta.period:
                seen.add(m.group(0))
    for old in seen:
        sep = "-" if "-" in old else ("~" if "~" in old else "–")
        rep(old, meta.period.replace("-", sep), "財年期間")
    return log


def update_asof_date(doc, as_of) -> list[str]:
    """製表／更新日期一律改成「這份報告實際產出的那天」。

    這種日期跟會議日期無關——範本裡 -3 是 7月23日、-4 是 7月24日、-6 是 7月25日，
    就是當初各檔更新的那天。`update_dates()` 的通用規則只換得到月份，會留下舊的
    「日」，於是報表上出現一個既不是會議日、也不是製表日的怪日期。

    只動兩種明確的位置，其餘一律不碰：
      ① 整段就是一個日期（-3 -4 -5 -6 -7 -8 -10 的標題下方）
      ② 「…2026年7月24日更新」（-2 的資料來源註記）
    像 `時間:2026年7月26日(星期日)` 或職員名單的 `2026年7月1日-2027年5月30日`
    都不符合這兩種形狀，不會被改到。
    """
    want = f"{as_of.year}年{as_of.month}月{as_of.day}日"
    log: list[str] = []
    for para in iter_paragraphs(doc):
        text = para.text.strip()
        if not text or "年" not in text:
            continue
        if _RE_YMD.fullmatch(text):
            if text != want:
                set_paragraph_text(para, want)
                log.append(f"製表日期: {text} → {want}")
            continue
        if "更新" in text:
            m = re.search(_RE_YMD.pattern + r"(?=\s*更新)", text)
            if m and m.group(0) != want:
                _replace_in_paragraph(para, m.group(0), want)
                log.append(f"資料更新日: {m.group(0)} → {want}")
    return log


def open_doc(path: str) -> Document:
    return Document(path)
